"""
RAG Ingestion Pipeline
Reads PDF, DOCX, MD, JSON files from models/sources/ and upserts them
into Qdrant using Gemini text-embedding-004.
"""
import os
import json
import glob
import logging
from pathlib import Path
from typing import List

from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from langchain_qdrant import QdrantVectorStore
from qdrant_client import QdrantClient
from qdrant_client.models import Distance, VectorParams

logger = logging.getLogger(__name__)

QDRANT_URL = os.getenv("QDRANT_URL", "http://localhost:6333")
COLLECTION_NAME = "agri_knowledge_base"
EMBEDDING_MODEL = "models/gemini-embedding-2"          # New Gemini embedding model
SOURCES_DIR = os.getenv("RAG_SOURCES_DIR", "models/sources")

# Resolve sources dir relative to repo root (3 levels up from this file)
_service_dir = Path(__file__).resolve().parent.parent.parent.parent.parent
SOURCES_PATH = _service_dir / SOURCES_DIR


def _get_embeddings() -> GoogleGenerativeAIEmbeddings:
    api_key = os.getenv("GEMINI_API_KEY")
    if not api_key:
        raise RuntimeError("GEMINI_API_KEY environment variable is not set.")
    return GoogleGenerativeAIEmbeddings(
        model=EMBEDDING_MODEL,
        google_api_key=api_key,
    )


def _ensure_collection(client: QdrantClient, vector_size: int = 3072):
    """Create Qdrant collection if it does not already exist."""
    existing = [c.name for c in client.get_collections().collections]
    if COLLECTION_NAME not in existing:
        client.create_collection(
            collection_name=COLLECTION_NAME,
            vectors_config=VectorParams(size=vector_size, distance=Distance.COSINE),
        )
        logger.info(f"Created Qdrant collection: {COLLECTION_NAME}")
    else:
        logger.info(f"Collection '{COLLECTION_NAME}' already exists — skipping creation.")


def _parse_yaml_frontmatter(text: str) -> tuple[dict, str]:
    """Parse YAML frontmatter from a markdown string."""
    metadata = {}
    content = text
    if text.strip().startswith("---"):
        parts = text.split("---", 2)
        if len(parts) >= 3:
            frontmatter_text = parts[1]
            content = parts[2]
            for line in frontmatter_text.splitlines():
                if ":" in line:
                    key, val = line.split(":", 1)
                    key = key.strip()
                    val = val.strip().strip('"').strip("'")
                    metadata[key] = val
    return metadata, content


def _load_documents() -> List[Document]:
    """Load all supported documents from the sources directory."""
    docs: List[Document] = []
    sources_path = Path(SOURCES_PATH)

    if not sources_path.exists():
        logger.warning(f"Sources directory not found: {sources_path}")
        return docs

    # ---------- JSON files ----------
    for fp in sources_path.glob("**/*.json"):
        try:
            with open(fp, "r", encoding="utf-8") as f:
                raw = json.load(f)
            # Flatten JSON into readable text chunks
            text = _flatten_json(raw)
            docs.append(Document(
                page_content=text,
                metadata={"source": fp.name, "type": "json"}
            ))
            logger.info(f"Loaded JSON: {fp.name}")
        except Exception as e:
            logger.error(f"Failed to load JSON {fp}: {e}")

    # ---------- Markdown files ----------
    for fp in sources_path.glob("**/*.md"):
        try:
            text = fp.read_text(encoding="utf-8")
            metadata, content = _parse_yaml_frontmatter(text)
            docs.append(Document(
                page_content=content,
                metadata={
                    "source": fp.name,
                    "type": "markdown",
                    "crop": metadata.get("crop", "Unknown"),
                    "disease_pest_name": metadata.get("disease_pest_name", metadata.get("topic", "Unknown")),
                    "publisher": metadata.get("publisher", "Unknown"),
                    "publication_year": metadata.get("year", metadata.get("publication_year", "Unknown")),
                    "doi": metadata.get("doi", "Unknown"),
                    "academic_citation": metadata.get("academic_citation", "Unknown"),
                }
            ))
            logger.info(f"Loaded Markdown with metadata: {fp.name}")
        except Exception as e:
            logger.error(f"Failed to load MD {fp}: {e}")

    # ---------- PDF files ----------
    try:
        from langchain_community.document_loaders import PyPDFLoader
        for fp in sources_path.glob("**/*.pdf"):
            try:
                loader = PyPDFLoader(str(fp))
                pages = loader.load()
                for page in pages:
                    page.metadata["source"] = fp.name
                    page.metadata["type"] = "pdf"
                docs.extend(pages)
                logger.info(f"Loaded PDF: {fp.name} ({len(pages)} pages)")
            except Exception as e:
                logger.error(f"Failed to load PDF {fp}: {e}")
    except ImportError:
        logger.warning("langchain_community not installed — PDF loading skipped.")

    # ---------- DOCX files ----------
    try:
        from docx import Document as DocxDocument
        for fp in sources_path.glob("**/*.docx"):
            try:
                docx = DocxDocument(str(fp))
                text = "\n".join([p.text for p in docx.paragraphs if p.text.strip()])
                docs.append(Document(
                    page_content=text,
                    metadata={"source": fp.name, "type": "docx"}
                ))
                logger.info(f"Loaded DOCX: {fp.name}")
            except Exception as e:
                logger.error(f"Failed to load DOCX {fp}: {e}")
    except ImportError:
        logger.warning("python-docx not installed — DOCX loading skipped.")

    return docs


def _flatten_json(obj, prefix="") -> str:
    """Recursively flatten a JSON object into readable key-value text."""
    parts = []
    if isinstance(obj, dict):
        for k, v in obj.items():
            key = f"{prefix}.{k}" if prefix else k
            parts.append(_flatten_json(v, key))
    elif isinstance(obj, list):
        for i, item in enumerate(obj):
            parts.append(_flatten_json(item, f"{prefix}[{i}]"))
    else:
        parts.append(f"{prefix}: {obj}")
    return "\n".join(parts)


def _split_documents(docs: List[Document]) -> List[Document]:
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200,
        separators=["\n\n", "\n", "।", ".", " ", ""],  # Bangla-aware separators
    )
    return splitter.split_documents(docs)


def run_ingestion() -> dict:
    """
    Main ingestion entry point.
    Loads all source documents, chunks them, embeds with Gemini, and upserts to Qdrant.
    Returns a status summary dict.
    """
    logger.info("=== Starting RAG Ingestion Pipeline ===")
    docs = _load_documents()
    if not docs:
        return {"status": "no_documents", "chunks_indexed": 0}

    chunks = _split_documents(docs)
    logger.info(f"Split into {len(chunks)} chunks.")

    embeddings = _get_embeddings()
    qdrant_api_key = os.getenv("QDRANT_API_KEY")

    # Initialize client explicitly with high timeout to prevent cloud timeouts
    client = QdrantClient(url=QDRANT_URL, api_key=qdrant_api_key, timeout=120)
    _ensure_collection(client)

    # Upload in batches of 32 to prevent write timeouts and rate limits
    batch_size = 32
    total_chunks = len(chunks)
    logger.info(f"Uploading {total_chunks} chunks to Qdrant in batches of {batch_size}...")

    # Initialize store using client
    store = QdrantVectorStore(
        client=client,
        collection_name=COLLECTION_NAME,
        embedding=embeddings,
    )

    # Upload batches with robust retry and rate limit handling
    import time
    for i in range(0, total_chunks, batch_size):
        batch = chunks[i : i + batch_size]
        
        max_retries = 7
        delay = 10
        for attempt in range(max_retries):
            try:
                store.add_documents(batch)
                logger.info(f"Uploaded chunks {i+1}-{min(i+batch_size, total_chunks)} of {total_chunks}")
                time.sleep(2.0)  # Pause to avoid rate limits
                break
            except Exception as e:
                err_str = str(e)
                if "429" in err_str or "RESOURCE_EXHAUSTED" in err_str or "Quota exceeded" in err_str:
                    logger.warning(f"Rate limit hit during batch upload. Retrying in {delay} seconds (Attempt {attempt+1}/{max_retries})...")
                    time.sleep(delay)
                    delay *= 2
                else:
                    raise e
        else:
            raise RuntimeError(f"Failed to upload batch starting at chunk {i} after {max_retries} retries.")


    logger.info(f"Upserted all {total_chunks} chunks to Qdrant collection '{COLLECTION_NAME}'.")
    return {
        "status": "success",
        "documents_loaded": len(docs),
        "chunks_indexed": len(chunks),
        "collection": COLLECTION_NAME,
    }
